import docx
# 在word中添加自定义样式
# 读取邀请名单
with open('guests.txt', 'r') as f:
    guests = f.readlines()
# 生成定制邀请函的 Word 文档
doc = docx.Document('guests.docx')
for guest in guests:
    doc.add_paragraph('It would be a pleasure to have the company of')\
        .style = 'Customsytleitalic'
    doc.add_paragraph(guest.strip())\
        .style = 'Customstyle'
    doc.add_paragraph('at 11010 Memory Lane on the Evening of')\
        .style = 'Customsytleitalic'
    doc.add_paragraph('April 1st')\
        .style = 'Customstyle'
    paraObj = doc.add_paragraph('at 7 o\'clock')
    paraObj.style = 'Customsytleitalic'
    # 每份邀请函应该占据一页
    doc.add_page_break()
doc.save('guests.docx')
